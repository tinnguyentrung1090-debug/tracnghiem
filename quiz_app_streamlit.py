import streamlit as st
import io
# Cần cài đặt thư viện python-docx: pip install python-docx
try:
    from docx import Document
except ImportError:
    st.error("Cần cài đặt thư viện 'python-docx' để đọc file Word: pip install python-docx")
    Document = None # Đặt Document = None nếu import thất bại
    
from google import genai
from google.genai.errors import APIError
import json

# --- Cấu hình Trang Streamlit ---
st.set_page_config(
    page_title="App Ôn Luyện Đề Thi Trắc Nghiệm",
    layout="wide"
)

st.title("📚 Ứng dụng Ôn Luyện Trắc Nghiệm (Word File)")

# Khóa API để gọi Gemini - Lấy từ Streamlit Secrets
API_KEY = st.secrets.get("GEMINI_API_KEY") 

# --- HÀM XỬ LÝ DOCX VÀ PARSING BẰNG GEMINI ---

@st.cache_data(show_spinner=False)
def read_docx(uploaded_file):
    """Đọc toàn bộ văn bản từ file Word (.docx)."""
    if Document is None:
        st.error("Không thể đọc file Word. Vui lòng kiểm tra đã cài đặt 'python-docx' chưa.")
        return None
    try:
        # docx.Document() có thể nhận file object (uploaded_file)
        document = Document(uploaded_file)
        full_text = []
        for para in document.paragraphs:
            # Loại bỏ các dòng trống
            if para.text.strip():
                 full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Lỗi đọc file Word: {e}")
        return None

@st.cache_data(show_spinner="Đang phân tích cấu trúc đề thi với AI...")
def parse_quiz_data_with_gemini(raw_text, api_key):
    """
    Sử dụng Gemini API để phân tích cú pháp (parse) văn bản thô thành cấu trúc JSON.
    Cần đảm bảo file Word có cấu trúc rõ ràng (ví dụ: Câu 1, A, B, C, D, Đáp án đúng là X).
    """
    if not api_key:
        st.error("Lỗi: Không tìm thấy Khóa API. Vui lòng cấu hình Khóa 'GEMINI_API_KEY' trong Streamlit Secrets.")
        return None
    
    if not raw_text:
        return []

    system_instruction = (
        "Bạn là một chuyên gia phân tích cú pháp đề thi trắc nghiệm. Nhiệm vụ của bạn là đọc văn bản thô "
        "từ file Word và trích xuất tất cả các câu hỏi trắc nghiệm thành một mảng JSON. "
        "Mỗi câu hỏi phải bao gồm 'question', 'options' (một mảng chứa tất cả các lựa chọn), "
        "và 'correct_answer' (chỉ là chữ cái (A, B, C...) hoặc nội dung văn bản đáp án đúng)."
        "Đảm bảo kết quả là JSON hợp lệ và CHỈ chứa JSON, không thêm lời nói đầu hay giải thích."
        "Bạn phải luôn luôn trả về một mảng JSON."
    )
    
    user_prompt = f"Trích xuất tất cả câu hỏi trắc nghiệm từ văn bản sau:\n\n---\n{raw_text}\n---"
    
    try:
        client = genai.Client(api_key=api_key)
        
        # Định nghĩa Schema cho kết quả JSON
        response_schema = {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "question": {"type": "STRING", "description": "Nội dung câu hỏi."},
                    "options": {
                        "type": "ARRAY",
                        "items": {"type": "STRING"},
                        "description": "Mảng chứa các lựa chọn đáp án (A, B, C...)."
                    },
                    "correct_answer": {"type": "STRING", "description": "Chữ cái (A, B, C...) hoặc nội dung của đáp án đúng."}
                },
                "required": ["question", "options", "correct_answer"]
            }
        }
        
        config = {
            "response_mime_type": "application/json",
            "response_schema": response_schema
        }

        response = client.models.generate_content(
            model='gemini-2.5-flash-preview-05-20',
            contents=user_prompt,
            config=config,
            system_instruction=system_instruction
        )
        
        json_data = response.text
        return json.loads(json_data)
        
    except APIError as e:
        st.error(f"Lỗi gọi Gemini API: {e}. Vui lòng kiểm tra Khóa API hoặc giới hạn sử dụng.")
        return None
    except json.JSONDecodeError:
        st.error("Lỗi phân tích cú pháp JSON từ API. Vui lòng kiểm tra lại cấu trúc đề thi trong file Word.")
        return None
    except Exception as e:
        st.error(f"Đã xảy ra lỗi không xác định trong quá trình phân tích: {e}")
        return None

# --- KHỞI TẠO VÀ QUẢN LÝ SESSION STATE ---

def initialize_session_state():
    """Khởi tạo các biến trạng thái cần thiết cho ứng dụng."""
    if 'quiz_data' not in st.session_state:
        st.session_state.quiz_data = [] # Dữ liệu đề thi đã parse
    if 'current_mode' not in st.session_state:
        st.session_state.current_mode = 'upload' # upload | menu | study | exam | review | result
    if 'current_index' not in st.session_state:
        st.session_state.current_index = 0
    if 'exam_answers' not in st.session_state:
        st.session_state.exam_answers = {} # {question_index: user_answer_text}
    if 'score' not in st.session_state:
        st.session_state.score = None # None | {'correct': 5, 'wrong': 3, 'review_q': [...]}

initialize_session_state()

# --- HÀM THIẾT LẬP CHẾ ĐỘ ---

def set_mode(mode):
    """Thiết lập chế độ, reset trạng thái nếu cần."""
    if mode == 'study' or mode == 'menu':
        st.session_state.current_index = 0
        st.session_state.current_mode = mode
        st.session_state.score = None
        # Xóa trạng thái chọn của chế độ study cũ
        keys_to_delete = [k for k in st.session_state.keys() if k.startswith('study_selected_')]
        for k in keys_to_delete:
            del st.session_state[k]

    elif mode == 'exam':
        st.session_state.current_index = 0
        st.session_state.exam_answers = {}
        st.session_state.current_mode = 'exam'
        st.session_state.score = None
    elif mode == 'review':
        # Chế độ ôn lại câu sai, sử dụng danh sách câu sai từ kết quả
        if st.session_state.score and st.session_state.score.get('review_q'):
            st.session_state.current_index = 0
            st.session_state.current_mode = 'review'
        else:
            set_mode('result') # Quay lại màn hình kết quả

# --- HÀM TIỆN ÍCH CHO ĐÁP ÁN ---

def get_question_data(q_index, mode):
    """Trả về dữ liệu câu hỏi và chỉ mục gốc (nếu đang ở chế độ review)."""
    if mode == 'review':
        if not st.session_state.score or not st.session_state.score.get('review_q'):
             return None, None
        
        original_q_index = st.session_state.score['review_q'][q_index]
        question_data = st.session_state.quiz_data[original_q_index]
        return question_data, original_q_index
    else:
        question_data = st.session_state.quiz_data[q_index]
        return question_data, q_index

def get_correct_answer_text(question_data):
    """Tìm đáp án đúng trong list options dựa trên correct_answer (chữ cái/text) và định dạng."""
    correct_key = question_data['correct_answer'].strip()
    options = question_data['options']
    option_labels = ['A', 'B', 'C', 'D', 'E', 'F']
    
    # Ưu tiên tìm theo chữ cái
    if len(correct_key) == 1 and correct_key.upper() in option_labels:
        idx = option_labels.index(correct_key.upper())
        if idx < len(options):
            return f"{option_labels[idx]}. {options[idx]}"
    
    # Fallback: Tìm theo nội dung văn bản
    for i, opt in enumerate(options):
        # So sánh nội dung văn bản
        if opt.strip() == correct_key.strip():
            return f"{option_labels[i]}. {opt}"
            
    # Fallback cuối cùng: Trả về văn bản gốc của đáp án đúng (có thể là chữ cái hoặc nội dung)
    return correct_key

# --- HÀM RENDER CÂU HỎI CHUNG ---

def render_question(q_index, mode):
    """Hiển thị một câu hỏi dựa trên chỉ mục và chế độ."""
    
    question_data, original_q_index = get_question_data(q_index, mode)
    
    if question_data is None:
         st.error("Không tìm thấy dữ liệu câu hỏi.")
         return
         
    # Tính toán tổng số câu hỏi cần hiển thị
    total_questions = len(st.session_state.quiz_data) if mode == 'exam' or mode == 'study' else len(st.session_state.score['review_q'])
    
    display_q_number = original_q_index + 1 if mode == 'review' else q_index + 1
    
    st.markdown(f"**Câu {display_q_number} / {total_questions}**")
    st.markdown(f"#### {question_data['question']}")

    # Tạo key duy nhất cho radio button
    radio_key = f"{mode}_q_{original_q_index}" if mode == 'review' else f"{mode}_q_{q_index}"
    
    # Chuẩn bị options để hiển thị (thêm chữ cái A, B, C...)
    labeled_options = []
    option_labels = ['A', 'B', 'C', 'D', 'E', 'F']
    for i, opt in enumerate(question_data['options']):
        if i < len(option_labels):
            labeled_options.append(f"{option_labels[i]}. {opt}")
        else:
            labeled_options.append(opt)

    # --- CHẾ ĐỘ ÔN LUYỆN (STUDY / REVIEW) ---
    if mode == 'study' or mode == 'review':
        
        # Callback khi chọn đáp án
        def handle_study_selection(selected_text):
            # Lưu đáp án đã chọn vào session state
            st.session_state[f"study_selected_{radio_key}"] = selected_text
            # Ghi nhận kết quả
            st.session_state.is_answered = True

        # Lấy đáp án đã chọn trong chế độ học tập (dùng key riêng)
        study_selected = st.session_state.get(f"study_selected_{radio_key}")

        # Radio button
        selected_option_text = st.radio(
            "Chọn đáp án:",
            options=labeled_options,
            key=radio_key,
            on_change=handle_study_selection,
            args=(st.session_state[radio_key],)
        )
        
        # Sau khi chọn, hiển thị kết quả
        if study_selected:
            correct_answer_text = get_correct_answer_text(question_data)
            
            # Kiểm tra xem đáp án chọn có chứa văn bản của đáp án đúng không
            is_correct = correct_answer_text in study_selected
            
            # Hiển thị kết quả
            if is_correct:
                st.success("✅ Chính xác! Bạn đã chọn đúng.")
            else:
                st.error("❌ Sai rồi.")
            
            # Hiển thị đáp án đúng
            st.info(f"Đáp án đúng là: **{correct_answer_text}**")
            
            st.markdown("---")
            
            # Nút Next
            if q_index < total_questions - 1:
                if st.button("Câu tiếp theo >>", key=f"next_study_{q_index}"):
                    st.session_state.current_index += 1
                    # Xóa trạng thái chọn của câu hiện tại để radio button reset
                    if f"study_selected_{radio_key}" in st.session_state:
                         del st.session_state[f"study_selected_{radio_key}"]
                    st.rerun()
            else:
                st.success(f"Chúc mừng, bạn đã hoàn thành phần {'ôn lại' if mode == 'review' else 'ôn luyện'}!")
                if st.button("Quay lại Menu Chính"):
                     set_mode('menu')
                     st.rerun()

    # --- CHẾ ĐỘ KIỂM TRA (EXAM) ---
    elif mode == 'exam':
        
        # Lấy đáp án đã chọn trước đó (nếu có)
        initial_value = st.session_state.exam_answers.get(q_index)
        
        # Tìm index của đáp án đã chọn (None nếu chưa chọn)
        initial_index = labeled_options.index(initial_value) if initial_value in labeled_options else None
        
        def handle_exam_selection():
             # Lưu đáp án vào session state ngay khi người dùng tương tác
             selected_text = st.session_state[radio_key]
             st.session_state.exam_answers[q_index] = selected_text

        st.radio(
            "Chọn đáp án:",
            options=labeled_options,
            key=radio_key,
            index=initial_index,
            on_change=handle_exam_selection
        )
        
        st.markdown("---")
        
        col_prev, col_next = st.columns([1, 1])
        
        with col_prev:
            if q_index > 0:
                if st.button("<< Câu trước", key="prev_exam"):
                    st.session_state.current_index -= 1
                    st.rerun()
        
        with col_next:
            if q_index < total_questions - 1:
                if st.button("Câu tiếp theo >>", key="next_exam"):
                    st.session_state.current_index += 1
                    st.rerun()
            elif q_index == total_questions - 1:
                # Nút nộp bài ở câu cuối
                if st.button("NỘP BÀI KIỂM TRA", use_container_width=True, type="primary"):
                    calculate_score()
                    st.rerun()

def calculate_score():
    """Tính toán điểm số và lưu vào session_state."""
    total = len(st.session_state.quiz_data)
    correct_count = 0
    wrong_count = 0
    review_q_indices = [] # Chỉ mục các câu hỏi sai (trong mảng quiz_data)
    
    for i, question_data in enumerate(st.session_state.quiz_data):
        user_answer_text = st.session_state.exam_answers.get(i)
        
        # Lấy văn bản đáp án đúng đã được định dạng (VD: 'B. Nội dung đáp án B')
        correct_answer_text = get_correct_answer_text(question_data)
        
        # Kiểm tra nếu người dùng đã chọn và đáp án đúng có trong nội dung đáp án chọn
        is_correct = user_answer_text and correct_answer_text in user_answer_text
        
        if is_correct:
            correct_count += 1
        else:
            wrong_count += 1
            # Chỉ thêm vào danh sách ôn lại nếu đã có câu trả lời (hoặc chưa trả lời, tính là sai)
            review_q_indices.append(i) 

    st.session_state.score = {
        'correct': correct_count,
        'wrong': wrong_count,
        'total': total,
        'review_q': review_q_indices
    }
    st.session_state.current_mode = 'result'

# --- HÀM RENDER CÁC MÀN HÌNH ---

def render_upload_screen():
    """Màn hình tải file và menu chính."""
    uploaded_file = st.file_uploader(
        "1. Tải file Word (.docx) chứa đề thi trắc nghiệm của bạn.",
        type=['docx']
    )
    
    if uploaded_file is not None and st.session_state.current_mode == 'upload':
        raw_text = read_docx(uploaded_file)
        if raw_text:
             if API_KEY:
                 parsed_data = parse_quiz_data_with_gemini(raw_text, API_KEY)
                 
                 if parsed_data and isinstance(parsed_data, list) and len(parsed_data) > 0:
                     st.session_state.quiz_data = parsed_data
                     st.success(f"Phân tích thành công! Tìm thấy {len(parsed_data)} câu hỏi.")
                     set_mode('menu')
                     st.rerun()
                 elif parsed_data is not None:
                      st.error("Không tìm thấy câu hỏi nào hoặc cấu trúc đề thi không rõ ràng. Vui lòng kiểm tra lại file Word.")
             else:
                 st.error("⚠️ Không tìm thấy API Key. Vui lòng thêm 'GEMINI_API_KEY' vào Streamlit Secrets.")

    if st.session_state.current_mode == 'menu' and st.session_state.quiz_data:
        render_menu_screen()
        
def render_menu_screen():
    """Màn hình chọn chế độ ôn luyện/kiểm tra."""
    st.header("2. Chọn Chế độ")
    st.subheader(f"Đã tải {len(st.session_state.quiz_data)} câu hỏi.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📝 Ôn Luyện (Study Mode)")
        st.info("**Tính năng:** Chọn đáp án và **hiện ngay kết quả đúng/sai**.")
        if st.button("BẮT ĐẦU ÔN LUYỆN", use_container_width=True, type="secondary"):
            set_mode('study')
            st.rerun()

    with col2:
        st.markdown("### ⏱️ Kiểm Tra (Exam Mode)")
        st.info("**Tính năng:** Làm bài ẩn, **nộp bài** để xem kết quả tổng quát.")
        if st.button("BẮT ĐẦU KIỂM TRA", use_container_width=True, type="primary"):
            set_mode('exam')
            st.rerun()
            
def render_quiz_main():
    """Hiển thị giao diện chính cho Study, Exam, hoặc Review."""
    
    mode = st.session_state.current_mode
    
    # Hiển thị nút quay lại Menu Chính ở trên cùng
    if st.button("<< Quay lại Menu Chính", key="main_back_to_menu"):
        set_mode('menu')
        st.rerun()
        
    st.markdown("---")
        
    if mode == 'study':
        st.header("📝 CHẾ ĐỘ ÔN LUYỆN")
    elif mode == 'exam':
        st.header("⏱️ CHẾ ĐỘ KIỂM TRA")
    elif mode == 'review':
        st.header("🔁 ÔN LẠI CÂU SAI")
    
    data_to_display = st.session_state.quiz_data
    if mode == 'review':
         # Đảm bảo list review_q tồn tại và không rỗng
         if not st.session_state.score or not st.session_state.score.get('review_q'):
              st.error("Không có câu hỏi sai để ôn lại.")
              set_mode('result')
              st.rerun()
              return
         data_to_display = st.session_state.score['review_q'] # Chỉ là index của câu hỏi sai
         
    data_len = len(data_to_display)

    if data_len > 0 and st.session_state.current_index < data_len:
        render_question(st.session_state.current_index, mode)
    elif data_len > 0:
        # Trường hợp đã hoàn thành (Study/Review)
        st.success(f"Bạn đã hoàn thành toàn bộ phần {'ôn lại' if mode == 'review' else 'ôn luyện'}!")
        if mode == 'review' and st.button("Quay lại Kết quả"):
             st.session_state.current_mode = 'result'
             st.rerun()

def render_result_screen():
    """Màn hình kết quả sau khi nộp bài."""
    score = st.session_state.score
    if not score:
        st.warning("Không có kết quả để hiển thị.")
        set_mode('menu')
        return

    st.header("🎉 KẾT QUẢ KIỂM TRA")
    st.markdown("---")
    
    # Tính điểm và tỷ lệ phần trăm
    percent_correct = score['correct'] / score['total'] * 100
    
    if percent_correct >= 80:
         st.balloons()
         st.success(f"Tuyệt vời! Bạn đạt {percent_correct:.1f}% số câu đúng.")
    elif percent_correct >= 50:
         st.info(f"Kết quả tốt! Bạn đạt {percent_correct:.1f}% số câu đúng.")
    else:
         st.warning(f"Cần cố gắng hơn! Bạn đạt {percent_correct:.1f}% số câu đúng.")


    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Tổng số câu", score['total'])
    with col2:
        st.metric("Số câu ĐÚNG", score['correct'])
    with col3:
        st.metric("Số câu SAI", score['wrong'])
        
    st.markdown("---")
    
    if score['wrong'] > 0:
        st.info(f"Hãy ôn luyện lại **{score['wrong']} câu hỏi** sai để củng cố kiến thức.")
        if st.button("🔁 ÔN LUYỆN LẠI CÂU SAI", type="primary", use_container_width=True):
            set_mode('review')
            st.rerun()
    else:
        st.success("Bạn đã trả lời đúng tất cả các câu hỏi!")
        
    if st.button("Quay lại Menu Chính", key="back_to_menu_from_result"):
        set_mode('menu')
        st.rerun()

# --- MAIN RENDER LOGIC ---

# Sidebar Navigation/Status
st.sidebar.title("Trạng thái và Cấu hình")
if st.session_state.quiz_data:
    st.sidebar.markdown(f"**Tổng câu hỏi:** **{len(st.session_state.quiz_data)}**")
    st.sidebar.markdown(f"**Chế độ hiện tại:** {st.session_state.current_mode.capitalize()}")

    if st.session_state.current_mode not in ['menu', 'upload']:
        st.sidebar.markdown(f"**Vị trí:** Câu **{st.session_state.current_index + 1}**")

# Hiển thị các màn hình dựa trên mode
if st.session_state.current_mode == 'upload' or st.session_state.current_mode == 'menu':
    render_upload_screen()
elif st.session_state.current_mode == 'result':
    render_result_screen()
elif st.session_state.current_mode in ['study', 'exam', 'review']:
    render_quiz_main()

# Lưu ý quan trọng cho người dùng về API Key và thư viện
st.sidebar.markdown("---")
st.sidebar.markdown("💡 **Lưu ý:**")
st.sidebar.markdown("- Ứng dụng cần thư viện `python-docx` (`pip install python-docx`).")
st.sidebar.markdown("- Cần cấu hình Khóa API **'GEMINI_API_KEY'** trong Streamlit Secrets để phân tích file Word.")
if not API_KEY:
     st.sidebar.error("⚠️ THIẾU API KEY! Vui lòng cấu hình ngay.")
